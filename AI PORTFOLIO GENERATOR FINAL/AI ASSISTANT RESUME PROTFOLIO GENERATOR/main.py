"""
AI-Assisted Resume Portfolio Generator - Perfection Engine
------------------------------------------------------------
Extracts structured JSON data & ATS score from resumes using Gemini 2.5 Flash, 
supports 4 visual custom themes, profile picture avatars, and interactive AI Chat.
"""

import os
import sys
import json
import re
import io
from pathlib import Path
from dotenv import load_dotenv

# Ensure UTF-8 output encoding for Windows terminals
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

# Import official Google GenAI SDK
try:
    from google import genai
    from google.genai import types
    SDK_TYPE = "genai"
except ImportError:
    try:
        import google.generativeai as legacy_genai
        SDK_TYPE = "legacy"
    except ImportError:
        SDK_TYPE = None

# Configuration Constants
RESUME_FILE = "resume.txt"
TEMPLATE_FILE = "template.html"
OUTPUT_FILE = "portfolio.html"
MIN_RESUME_LENGTH = 30


def step_logger(step_num: int, message: str):
    """Prints step indicators to give friendly console feedback."""
    print(f"\n[Step {step_num}/5] {message}")


def load_api_key() -> str:
    """Loads environment variables from .env file and validates GEMINI_API_KEY."""
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    
    if not api_key or api_key.strip() == "" or api_key == "your_gemini_api_key_here":
        print("\n" + "=" * 70)
        print("❌ ERROR: GEMINI_API_KEY is missing or invalid in your .env file!")
        print("=" * 70)
        sys.exit(1)
        
    return api_key.strip()


def extract_text_from_file(file_path_or_bytes, filename: str) -> str:
    """Extracts text content from TXT, PDF, DOCX, or DOC files."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        try:
            import pypdf
            if isinstance(file_path_or_bytes, (str, Path)):
                reader = pypdf.PdfReader(file_path_or_bytes)
            else:
                reader = pypdf.PdfReader(io.BytesIO(file_path_or_bytes))
            
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            return text.strip()
        except Exception as e:
            print(f"⚠️ pypdf fallback: {e}")
            return ""

    elif ext in [".docx", ".doc"]:
        try:
            import docx
            if isinstance(file_path_or_bytes, (str, Path)):
                doc = docx.Document(file_path_or_bytes)
            else:
                doc = docx.Document(io.BytesIO(file_path_or_bytes))
            
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return text.strip()
        except Exception as e:
            print(f"⚠️ docx fallback: {e}")
            return ""

    else:
        if isinstance(file_path_or_bytes, (str, Path)):
            with open(file_path_or_bytes, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        else:
            return file_path_or_bytes.decode("utf-8", errors="ignore").strip()


def read_and_validate_resume(filepath: str) -> str:
    """Reads resume file and validates minimum content length."""
    path = Path(filepath)
    if not path.exists():
        print(f"\n❌ ERROR: Resume file '{filepath}' was not found!")
        sys.exit(1)

    content = extract_text_from_file(path, path.name)

    if len(content) < MIN_RESUME_LENGTH:
        print(f"\n❌ ERROR: '{filepath}' contains insufficient text ({len(content)} chars).")
        sys.exit(1)

    return content


def extract_data_with_gemini(resume_input, filename: str, api_key: str) -> str:
    """
    Sends resume to Gemini 2.5 Flash. Generates structured JSON + ATS Analysis.
    """
    if SDK_TYPE is None:
        print("\n❌ ERROR: 'google-genai' SDK is not installed.")
        sys.exit(1)

    system_instruction = (
        "You are an expert AI Resume Parser & ATS Advisor. Extract ALL information "
        "from the provided resume document into a complete structured JSON payload.\n\n"
        "STRICT EXTRACTION RULES:\n"
        "1. Extract ALL details: Name, Headline, Summary, Contact, Skills (categorized), Experience (with all bullets), Projects, Education, and Achievements.\n"
        "2. Calculate an 'ats_analysis' object with a quality score (0-100), key strengths, and 3 actionable improvement suggestions.\n"
        "3. Do NOT hallucinate. Return null or empty list [] for unmentioned fields.\n"
        "4. Return ONLY valid JSON matching the exact schema below.\n\n"
        "JSON SCHEMA:\n"
        "{\n"
        '  "name": "Full Name or null",\n'
        '  "headline": "Professional Title / Tagline or null",\n'
        '  "summary": "Professional Summary or null",\n'
        '  "contact": {\n'
        '    "email": "Email or null",\n'
        '    "phone": "Phone or null",\n'
        '    "location": "Location or null",\n'
        '    "linkedin": "LinkedIn URL or null",\n'
        '    "github": "GitHub URL or null",\n'
        '    "portfolio": "Website URL or null"\n'
        "  },\n"
        '  "ats_analysis": {\n'
        '    "score": 92,\n'
        '    "strengths": ["Clear technical section", "Quantified bullet points"],\n'
        '    "suggestions": ["Add more cloud deployment metrics", "Include specific metrics in project outcomes"]\n'
        "  },\n"
        '  "stats": {\n'
        '    "years_experience": "Years (e.g. 5+)",\n'
        '    "projects_count": "Number (e.g. 10+)",\n'
        '    "awards_count": "Number (e.g. 4+)"\n'
        "  },\n"
        '  "skills": [\n'
        '    {\n'
        '      "category": "Category Name",\n'
        '      "items": ["Skill 1", "Skill 2"]\n'
        '    }\n'
        "  ],\n"
        '  "experience": [\n'
        '    {\n'
        '      "title": "Job Title",\n'
        '      "company": "Company Name",\n'
        '      "period": "Dates",\n'
        '      "bullets": ["Bullet 1", "Bullet 2"]\n'
        '    }\n'
        "  ],\n"
        '  "projects": [\n'
        '    {\n'
        '      "title": "Project Name",\n'
        '      "tech_stack": ["Tech 1"],\n'
        '      "description": "Description",\n'
        '      "link": "URL or null"\n'
        '    }\n'
        "  ],\n"
        '  "education": [\n'
        '    {\n'
        '      "degree": "Degree",\n'
        '      "institution": "University / School",\n'
        '      "period": "Dates",\n'
        '      "details": "GPA/Details or null"\n'
        '    }\n'
        "  ],\n"
        '  "achievements": ["Achievement 1"]\n'
        "}"
    )

    ext = Path(filename).suffix.lower()

    try:
        if SDK_TYPE == "genai":
            client = genai.Client(api_key=api_key)

            if isinstance(resume_input, bytes) and ext in [".pdf", ".png", ".jpg", ".jpeg"]:
                mime_type = "application/pdf" if ext == ".pdf" else f"image/{ext.replace('.', '')}"
                part = types.Part.from_bytes(data=resume_input, mime_type=mime_type)
                contents = [part, system_instruction]
            else:
                text_content = resume_input.decode("utf-8", errors="ignore") if isinstance(resume_input, bytes) else str(resume_input)
                contents = [f"{system_instruction}\n\nRESUME CONTENT:\n{text_content}"]

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.1
                )
            )
            return response.text
        else:
            legacy_genai.configure(api_key=api_key)
            model = legacy_genai.GenerativeModel("gemini-1.5-flash")
            text_content = resume_input.decode("utf-8", errors="ignore") if isinstance(resume_input, bytes) else str(resume_input)
            response = model.generate_content(
                f"{system_instruction}\n\nRESUME CONTENT:\n{text_content}",
                generation_config={"response_mime_type": "application/json", "temperature": 0.1}
            )
            return response.text

    except Exception as e:
        print(f"❌ Gemini API Error: {e}")
        raise RuntimeError(f"Gemini API processing failed: {e}")


def parse_and_clean_json(raw_json_str: str) -> dict:
    """Safely parses JSON string from Gemini."""
    cleaned_str = raw_json_str.strip()
    
    if cleaned_str.startswith("```"):
        cleaned_str = re.sub(r"^```(?:json)?\s*", "", cleaned_str, flags=re.IGNORECASE)
        cleaned_str = re.sub(r"\s*```$", "", cleaned_str).strip()

    try:
        data = json.loads(cleaned_str)
    except json.JSONDecodeError as err:
        print(f"❌ JSON Decode Error: {err}")
        raise ValueError(f"Failed to parse AI output into valid JSON: {err}")

    default_schema = {
        "name": None, "headline": None, "summary": None,
        "contact": {}, "ats_analysis": {}, "stats": {}, "skills": [],
        "experience": [], "projects": [], "education": [], "achievements": []
    }
    
    for k, v in default_schema.items():
        if k not in data or data[k] is None:
            data[k] = v

    return data


def answer_recruiter_question(query: str, portfolio_data: dict, api_key: str) -> str:
    """Answers recruiter chatbot questions grounded in the user's portfolio data."""
    if not portfolio_data or not isinstance(portfolio_data, dict) or not portfolio_data.get("name"):
        # Fallback to reading resume.txt if available
        if Path(RESUME_FILE).exists():
            try:
                resume_text = read_and_validate_resume(RESUME_FILE)
                context_str = f"Resume Text Content:\n{resume_text}"
            except Exception:
                context_str = "No specific portfolio data loaded yet."
        else:
            context_str = "No specific portfolio data loaded yet."
    else:
        context_str = json.dumps(portfolio_data, indent=2)

    prompt = (
        "You are an AI Recruiter Assistant representing the candidate described in the data below.\n"
        "Answer the recruiter's question directly, politely, and professionally in 2-3 sentences max.\n"
        "Use ONLY facts present in the portfolio data. If asked something not in the data, state politely that it's not listed.\n\n"
        f"CANDIDATE PORTFOLIO DATA:\n{context_str}\n\n"
        f"RECRUITER QUESTION: {query}\n\n"
        "RESPONSE:"
    )

    try:
        client = genai.Client(api_key=api_key)
        resp = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.2)
        )
        return resp.text.strip()
    except Exception as e:
        return f"I'm sorry, I encountered an issue accessing candidate details: {e}"


def build_contact_pills(contact: dict) -> str:
    """Builds HTML pills for contact details."""
    if not isinstance(contact, dict):
        return ""
        
    pills = []
    
    email = contact.get("email")
    if email:
        pills.append(f'<a href="mailto:{email}" class="contact-pill"><span>✉</span> {email}</a>')

    phone = contact.get("phone")
    if phone:
        pills.append(f'<a href="tel:{phone}" class="contact-pill"><span>📞</span> {phone}</a>')

    location = contact.get("location")
    if location:
        pills.append(f'<div class="contact-pill"><span>📍</span> {location}</div>')

    linkedin = contact.get("linkedin")
    if linkedin:
        url = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
        pills.append(f'<a href="{url}" target="_blank" rel="noopener" class="contact-pill"><span>🔗</span> LinkedIn</a>')

    github = contact.get("github")
    if github:
        url = github if github.startswith("http") else f"https://{github}"
        pills.append(f'<a href="{url}" target="_blank" rel="noopener" class="contact-pill"><span>💻</span> GitHub</a>')

    portfolio = contact.get("portfolio")
    if portfolio:
        url = portfolio if portfolio.startswith("http") else f"https://{portfolio}"
        pills.append(f'<a href="{url}" target="_blank" rel="noopener" class="contact-pill"><span>🌐</span> Portfolio</a>')

    return "\n        ".join(pills)


def build_skills_html(skills: list) -> str:
    """Builds HTML grid cards for categorized skills."""
    if not isinstance(skills, list) or len(skills) == 0:
        return ""

    cards = []
    for skill_cat in skills:
        if isinstance(skill_cat, dict):
            category = skill_cat.get("category", "General Skills")
            items = skill_cat.get("items", [])
        elif isinstance(skill_cat, str):
            category = "Skills"
            items = [skill_cat]
        else:
            continue

        if not items:
            continue

        tag_spans = "".join([f'<span class="skill-pill">{item}</span>' for item in items])
        card_html = (
            f'<div class="skill-card">\n'
            f'  <div class="skill-category-name">{category}</div>\n'
            f'  <div class="skill-pill-tags">\n    {tag_spans}\n  </div>\n'
            f'</div>'
        )
        cards.append(card_html)

    return "\n".join(cards)


def build_experience_html(experience: list) -> str:
    """Builds timeline HTML steps for experience."""
    if not isinstance(experience, list) or len(experience) == 0:
        return ""

    items_html = []
    for job in experience:
        if not isinstance(job, dict):
            continue

        title = job.get("title", "Role")
        company = job.get("company", "")
        period = job.get("period", "")
        bullets = job.get("bullets", [])

        bullet_lis = "".join([f'<li>{b}</li>' for b in bullets])
        bullets_ul = f'<ul class="step-bullets">\n{bullet_lis}\n</ul>' if bullets else ''

        job_html = (
            f'<div class="timeline-step">\n'
            f'  <div class="step-node"></div>\n'
            f'  <div class="step-period">{period}</div>\n'
            f'  <div class="step-title">{title}</div>\n'
            f'  <div class="step-subtitle">{company}</div>\n'
            f'  {bullets_ul}\n'
            f'</div>'
        )
        items_html.append(job_html)

    return "\n".join(items_html)


def build_education_html(education: list) -> str:
    """Builds HTML timeline steps for education background."""
    if not isinstance(education, list) or len(education) == 0:
        return ""

    cards = []
    for edu in education:
        if not isinstance(edu, dict):
            continue

        degree = edu.get("degree", "Degree")
        institution = edu.get("institution", "")
        period = edu.get("period", "")
        details = edu.get("details", "")

        details_p = f'<p class="step-bullets">{details}</p>' if details else ''

        edu_html = (
            f'<div class="timeline-step">\n'
            f'  <div class="step-node"></div>\n'
            f'  <div class="step-period">{period}</div>\n'
            f'  <div class="step-title">{degree}</div>\n'
            f'  <div class="step-subtitle">{institution}</div>\n'
            f'  {details_p}\n'
            f'</div>'
        )
        cards.append(edu_html)

    return "\n".join(cards)


def build_projects_html(projects: list) -> str:
    """Builds HTML cards for projects."""
    if not isinstance(projects, list) or len(projects) == 0:
        return ""

    cards = []
    for proj in projects:
        if not isinstance(proj, dict):
            continue

        title = proj.get("title", "Project")
        description = proj.get("description", "")
        tech_stack = proj.get("tech_stack", [])
        link = proj.get("link")

        tech_tags = "".join([f'<span class="tech-tag">{t}</span>' for t in tech_stack])
        tech_div = f'<div class="project-tech">{tech_tags}</div>' if tech_tags else ''

        link_html = ''
        if link:
            url = link if link.startswith("http") else f"https://{link}"
            link_html = f'<a href="{url}" target="_blank" rel="noopener" class="project-link-btn">View Project &rarr;</a>'

        proj_html = (
            f'<div class="project-card">\n'
            f'  <div>\n'
            f'    <div class="project-title">{title}</div>\n'
            f'    <p class="project-desc">{description}</p>\n'
            f'  </div>\n'
            f'  <div>\n'
            f'    {tech_div}\n'
            f'    {link_html}\n'
            f'  </div>\n'
            f'</div>'
        )
        cards.append(proj_html)

    return "\n".join(cards)


def build_achievements_html(achievements: list) -> str:
    """Builds cards for achievements."""
    if not isinstance(achievements, list) or len(achievements) == 0:
        return ""

    items = []
    for ach in achievements:
        if not ach:
            continue
        item_html = (
            f'<div class="skill-card">\n'
            f'  <div class="skill-category-name">🏆 Achievement / Certification</div>\n'
            f'  <p style="color: var(--text-gray); font-size: 0.9rem;">{ach}</p>\n'
            f'</div>'
        )
        items.append(item_html)

    return "\n".join(items)


def generate_portfolio(data: dict, template_file: str, output_file: str, profile_img_b64: str = None, theme: str = "dark-lime"):
    """Injects data into template.html and writes portfolio.html with selected theme."""
    if not Path(template_file).exists():
        print(f"❌ ERROR: Template '{template_file}' missing!")
        sys.exit(1)

    with open(template_file, "r", encoding="utf-8") as f:
        html = f.read()

    name = data.get("name") or "Portfolio Owner"
    first_name = name.split()[0] if name else "Candidate"
    initials = "".join([part[0].upper() for part in name.split() if part])[:2] if name else "PO"
    headline = data.get("headline") or ""
    summary = data.get("summary") or ""

    stats = data.get("stats", {})
    years_exp = stats.get("years_experience") or f"{len(data.get('experience', []))}+ Years"
    projects_cnt = stats.get("projects_count") or f"{len(data.get('projects', []))}+"
    awards_cnt = stats.get("awards_count") or f"{len(data.get('achievements', []))}+"

    # Profile Avatar HTML
    if profile_img_b64:
        avatar_html = f'<img src="{profile_img_b64}" alt="{name}" class="hero-profile-img">'
    else:
        avatar_html = f'<div class="hero-avatar-large">{initials}</div>'

    contact_pills_html = build_contact_pills(data.get("contact", {}))
    skills_html = build_skills_html(data.get("skills", []))
    experience_html = build_experience_html(data.get("experience", []))
    projects_html = build_projects_html(data.get("projects", []))
    education_html = build_education_html(data.get("education", []))
    achievements_html = build_achievements_html(data.get("achievements", []))

    replacements = {
        "{{THEME}}": theme,
        "{{NAME}}": name,
        "{{FIRST_NAME}}": first_name,
        "{{INITIALS}}": initials,
        "{{PROFILE_AVATAR_HTML}}": avatar_html,
        "{{HEADLINE}}": headline,
        "{{SUMMARY}}": summary,
        "{{YEARS_EXPERIENCE}}": years_exp,
        "{{PROJECTS_COUNT}}": projects_cnt,
        "{{AWARDS_COUNT}}": awards_cnt,
        "{{CONTACT_PILLS}}": contact_pills_html,
        "{{SKILLS_GRID}}": skills_html,
        "{{EXPERIENCE_TIMELINE}}": experience_html,
        "{{PROJECTS_GRID}}": projects_html,
        "{{EDUCATION_GRID}}": education_html,
        "{{ACHIEVEMENTS_LIST}}": achievements_html,
        "{{SKILLS_HIDDEN}}": "section-hidden" if not skills_html else "",
        "{{EXPERIENCE_HIDDEN}}": "section-hidden" if (not experience_html and not education_html) else "",
        "{{PROJECTS_HIDDEN}}": "section-hidden" if not projects_html else "",
        "{{ACHIEVEMENTS_HIDDEN}}": "section-hidden" if not achievements_html else "",
    }

    for key, val in replacements.items():
        html = html.replace(key, val)

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(html)

    print(f"✅ Generated output saved to '{output_file}' with theme '{theme}'")


def main():
    print("=" * 70)
    print("🚀 AI-ASSISTED RESUME PORTFOLIO GENERATOR")
    print("=" * 70)

    step_logger(1, "Loading Environment Variables (.env)...")
    api_key = load_api_key()

    step_logger(2, f"Reading and validating '{RESUME_FILE}'...")
    resume_text = read_and_validate_resume(RESUME_FILE)

    step_logger(3, "Sending resume text to Google Gemini API...")
    raw_response = extract_data_with_gemini(resume_text, RESUME_FILE, api_key)

    step_logger(4, "Parsing & validating structured JSON data...")
    portfolio_data = parse_and_clean_json(raw_response)

    step_logger(5, f"Rendering '{OUTPUT_FILE}'...")
    generate_portfolio(portfolio_data, TEMPLATE_FILE, OUTPUT_FILE)

    print("\n" + "=" * 70)
    print("🎉 SUCCESS! Your HTML Portfolio is ready.")
    print("=" * 70 + "\n")


if __name__ == "__main__":
    main()
